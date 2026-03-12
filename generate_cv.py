#!/usr/bin/env python3
"""Generate a modern, subtle DOCX CV for Antoine Jaussoin, openable in Apple Pages."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

# ── Colour palette (subtle blue-grey) ──────────────────────────────
ACCENT = RGBColor(0x2B, 0x6C, 0xB3)       # muted steel-blue
DARK   = RGBColor(0x1A, 0x1A, 0x2E)       # near-black for headings
BODY   = RGBColor(0x33, 0x33, 0x33)       # dark grey for body
LIGHT  = RGBColor(0x77, 0x77, 0x77)       # light grey for secondary text
RULE   = RGBColor(0xCC, 0xCC, 0xCC)       # line colour

FONT_BODY   = "Helvetica Neue"
FONT_HEAD   = "Helvetica Neue"

# ── Helpers ─────────────────────────────────────────────────────────
def set_cell_border(cell, **kwargs):
    """Set cell border. Usage: set_cell_border(cell, top={"sz": 4, "color": "CCCCCC"})"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{attrs["sz"]}" w:space="0" w:color="{attrs["color"]}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_thin_rule(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_section_heading(doc, text):
    """Add a styled section heading with a thin rule underneath."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.font.name = FONT_HEAD
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT
    run.font.bold = True
    run.font.all_caps = True
    # thin underline
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="2" w:color="2B6CB3"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def strip_markdown(text):
    """Very basic markdown stripping: bold, links."""
    text = text.replace("  \n", "\n").strip()
    return text


def add_rich_text(paragraph, text, base_size=Pt(9.5), base_color=BODY):
    """Parse simple markdown (bold, links) into runs."""
    text = text.replace("  \n", "\n").replace("\n\n", "\n")
    # Split on **bold** and [text](url) patterns
    parts = re.split(r'(\*\*.*?\*\*|\[.*?\]\(.*?\))', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
            run.font.size = base_size
            run.font.color.rgb = base_color
            run.font.name = FONT_BODY
        elif re.match(r'\[.*?\]\(.*?\)', part):
            m = re.match(r'\[(.*?)\]\((.*?)\)', part)
            display = m.group(1)
            run = paragraph.add_run(display)
            run.font.size = base_size
            run.font.color.rgb = ACCENT
            run.font.name = FONT_BODY
            run.font.underline = True
        else:
            run = paragraph.add_run(part)
            run.font.size = base_size
            run.font.color.rgb = base_color
            run.font.name = FONT_BODY


def format_date(d):
    """Format YYYY-MM-DD to Mon YYYY."""
    from datetime import datetime
    try:
        dt = datetime.strptime(d, "%Y-%m-%d")
        return dt.strftime("%b %Y")
    except Exception:
        return d


# ── CV Data ─────────────────────────────────────────────────────────
work = [
    {
        "title": "Lead Engineer",
        "company": "Balyasny Asset Management",
        "location": "London",
        "type": "Permanent",
        "from": "2020-01-01",
        "to": None,
        "techs": ["React", "TypeScript", "Styled Components", "C#", "TanStack Query", "Docker / K8s"],
        "description": "As Balyasny Asset Management's Lead Engineer, working on various projects across the company.",
    },
    {
        "title": "UI Tech Lead",
        "company": "Wayve",
        "location": "London",
        "type": "Permanent",
        "from": "2020-01-01",
        "to": None,
        "techs": ["React", "TypeScript", "Styled Components", "Node", "Docker / K8s", "Python"],
        "description": "As Wayve's UI Tech lead, working on various projects across the company. The most challenging one is the in-car UI, delivering a high performance experience critical to safety. The UI receives about 300 websocket frames per second, and yet maintains a constant 60 fps.",
    },
    {
        "title": "Senior Front-End React Engineer",
        "company": "Two Sigma",
        "location": "London",
        "type": "Permanent",
        "from": "2017-03-01",
        "to": "2020-01-01",
        "techs": ["React", "TypeScript", "Styled Components", "Node.js"],
        "description": "Front-End engineer on Two Sigma's Venn platform. Venn makes it simpler to select managers, choose assets and manage risk, via Two Sigma's factors analysis. The front-end was originally written in JavaScript (ES6), React, Redux, Redux-Saga and SCSS (with CSS Modules), then rewritten to TypeScript, replacing Redux by React's context and local state, and SCSS/CSS Modules by Styled Components. Part of the original team of 2 Front-End developers, playing a decisive role on the original architecture and rewrite.",
    },
    {
        "title": "Front Office / Front-end React Engineer",
        "company": "Royal Bank of Scotland",
        "location": "London",
        "type": "Contractor / Permanent",
        "from": "2012-07-01",
        "to": "2017-03-01",
        "techs": ["React", "Redux", "AngularJS", "Node.js"],
        "description": "Front-End developer on RBS' flagship platform \"Agile Markets\". Implementation of several modules, from FX options tickets (Peg, TWAP, OCO, IDO…) to analysis tools. Modules developed using React, AngularJS, or plain JavaScript (with D3). Connected to back-end services using REST endpoints and websockets (Caplin), tested using Mocha or Karma against a Node server.",
    },
    {
        "title": "ASP.NET MVC Front-End Developer",
        "company": "Bank of America Merrill Lynch",
        "location": "London",
        "type": "Contractor",
        "from": "2011-04-01",
        "to": "2012-06-01",
        "techs": ["C#", "NHibernate", "ASP.NET MVC", "ExtJS", "jQuery", "SQL Server"],
        "description": "Maintenance of a risk application in C# and ASP.NET MVC using NHibernate. The application interacts with MSCI Risk Metrics for risk analysis. Over 5,000 unit tests written using NUnit, Moq, SpecFlow and Selenium.",
    },
    {
        "title": "ASP.NET MVC Front-end Developer",
        "company": "Royal Bank of Scotland",
        "location": "London",
        "type": "Contractor",
        "from": "2011-01-01",
        "to": "2011-03-01",
        "techs": ["C#", "NHibernate", "MVC 3", "ExtJS", "AutoFac", "SQL Server"],
        "description": "Design and implementation of a Human Resource system consolidating data from multiple existing systems within RBS. Front end based on ASP.NET MVC 3 (Razor) and ExtJS, back end in C# 4 with NHibernate.",
    },
    {
        "title": "Silverlight / C# Developer",
        "company": "Royal Bank of Scotland",
        "location": "London",
        "type": "Contractor",
        "from": "2010-03-01",
        "to": "2010-12-01",
        "techs": ["Silverlight", "C#", "NHibernate", "WCF", "Autofac", "SQL Server"],
        "description": "Design and implementation of a global technology platform for browsing content stored in legal and credit documentation. Silverlight-based client used around the globe to manage RBS contracts and netting calculations.",
    },
    {
        "title": "Front Office Developer, Derivatives",
        "company": "Credit Suisse",
        "location": "London",
        "type": "Full-Time, AVP",
        "from": "2009-11-01",
        "to": "2010-03-01",
        "techs": ["ASP.NET", "C#", "SQL Server"],
        "description": "Design, maintenance and improvements of a derivatives trading platform, used both internally and externally by Credit Suisse customers.",
    },
    {
        "title": "Full-Stack Developer, Architect",
        "company": "BNP Paribas - FundQuest UK",
        "location": "London",
        "type": "Full-Time",
        "from": "2007-01-01",
        "to": "2009-11-01",
        "techs": ["ASP.NET", "C#", "NHibernate", "MySQL"],
        "description": "Design, implementation and maintenance of a Quantitative, Document Management and Trading Web Application. Complete architectural redesign: quantitative analysis tools, document management with full-text search, and a trading system linked to electronic trading platforms.",
    },
    {
        "title": ".NET Developer",
        "company": "British Telecom PLC",
        "location": "London",
        "type": "Full-Time",
        "from": "2006-09-01",
        "to": "2007-01-01",
        "techs": ["ASP.NET", "C#", "VB.NET", "DotNetNuke"],
        "description": "Development of a portal (sdk.bt.com) in ASP.NET using DotNetNuke. Agile (Scrum) methodology with 3-month cycles and 2-week iterations, using continuous integration (CruiseControl.NET).",
    },
    {
        "title": "Full-Stack Developer",
        "company": "Everydev",
        "location": "Évreux (France)",
        "type": "Associate & Co-founder",
        "from": "2004-12-01",
        "to": "2008-07-01",
        "techs": ["C#", "WinForm", "ASP.NET", "MySQL"],
        "description": "Co-founded a computer services company. Developed an online and offline backup software (FaciloSave) in C# featuring encryption, compression, custom networking protocols, and localization. Also developed the company website in ASP.NET with NHibernate.",
    },
    {
        "title": "Software Developer",
        "company": "Intuition Informatique",
        "location": "Évreux (France)",
        "type": "Internship",
        "from": "2003-11-01",
        "to": "2004-12-01",
        "techs": ["C#", "MySQL"],
        "description": "Part-time role (14h/week) as a programmer, web designer, and system administrator. Built a management software for assistance contracts in C#, integrating with an ERP system.",
    },
]

education = [
    {
        "school": "Supinfo Paris - Oxford Brookes University",
        "diploma": "Master (MSc) in Computer Science",
        "location": "Paris (2003-2005), Oxford (2005-2006)",
        "date": "2006-06-01",
        "description": "Three-year engineering school resulting in a European Master in Computer Science. Final year at Oxford Brookes University.",
    },
    {
        "school": "Lycée Richelieu - CPGE TSI",
        "diploma": "Cours Préparatoires aux Grandes Ecoles",
        "location": "Rueil-Malmaison (France)",
        "date": "2003-06-01",
        "description": "Math sup/Math Spé: intensive preparation in physics and mathematics for the Grandes Ecoles.",
    },
    {
        "school": "Passy-Buzenval",
        "diploma": "BAC STI (French A-Level) - With Honours",
        "location": "Rueil-Malmaison (France)",
        "date": "2001-07-01",
        "description": "",
    },
]

skills = [
    {"name": "TypeScript", "level": "Expert", "experience": 3,
     "related": ["ESLint", "Jest", "Create-React-App"]},
    {"name": "JavaScript", "level": "Expert", "experience": 7,
     "related": ["React", "Lodash", "date-fns", "Jest", "Prettier"]},
    {"name": "React", "level": "Expert", "experience": 5,
     "related": ["Redux", "Redux-Saga", "Socket.IO", "Webpack", "Hooks", "Recoil.js"]},
    {"name": "HTML & CSS", "level": "Expert", "experience": 7,
     "related": ["CSS Modules", "SCSS/SASS", "Styled Components"]},
    {"name": "Node", "level": "Advanced", "experience": 4,
     "related": ["Create-React-App", "Express"]},
    {"name": ".NET", "level": "Expert", "experience": 8,
     "related": ["C#", "ASP.NET MVC", "NHibernate", "NUnit", "SQL Server"]},
]

projects = [
    {
        "name": "Retrospected",
        "description": "Agile Retrospective Board using React, Redux and Socket.IO. Open-source project available at www.retrospected.com and github.com/antoinejaussoin/retro-board.",
    },
    {
        "name": "React VR Player",
        "description": "360° Virtual Reality video player component for the Oculus Rift, built with React.",
    },
]

profile = (
    "I'm a Senior Full-Stack Engineer, with strong experience in the financial industry. "
    "I am currently working at Balyasny Asset Management as the Lead Engineer. "
    "I previously worked at Wayve, an autonomous driving company, as the UI Tech Lead, "
    "delivering high-performance in-car UIs. Before that I was at Two Sigma, a systematic "
    "trading hedge-fund based in New York. "
    "I've had the opportunity to work on various greenfield projects in my career, from the "
    "Cudos platform at BNP (still live 10 years after!) to creating my own software company. "
    "I'm also the author of a few open-source projects, including an Agile Retrospective board "
    "(www.retrospected.com)."
)


# ── Build Document ──────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# Remove default spacing from Normal style
style = doc.styles["Normal"]
style.font.name  = FONT_BODY
style.font.size  = Pt(9.5)
style.font.color.rgb = BODY
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after  = Pt(2)
style.paragraph_format.line_spacing = Pt(13)

# ── Header ─────────────────────────────────────────────────────────────
p_name = doc.add_paragraph()
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_name.paragraph_format.space_after = Pt(0)
run = p_name.add_run("ANTOINE JAUSSOIN")
run.font.name = FONT_HEAD
run.font.size = Pt(24)
run.font.color.rgb = DARK
run.font.bold = False

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(2)
p_title.paragraph_format.space_after = Pt(0)
run = p_title.add_run("Senior Full Stack Engineer")
run.font.name = FONT_HEAD
run.font.size = Pt(13)
run.font.color.rgb = ACCENT
run.font.bold = False

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(2)
p_sub.paragraph_format.space_after = Pt(2)
run = p_sub.add_run("React  ·  TypeScript/JavaScript  ·  Node  ·  .NET")
run.font.name = FONT_HEAD
run.font.size = Pt(9.5)
run.font.color.rgb = LIGHT

# Contact line
p_contact = doc.add_paragraph()
p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_contact.paragraph_format.space_before = Pt(4)
p_contact.paragraph_format.space_after = Pt(4)
contact_items = [
    "antoine@jaussoin.com",
    "+44 77 22 55 77 39",
    "www.jaussoin.com",
    "London SW11 5QA",
]
for i, item in enumerate(contact_items):
    run = p_contact.add_run(item)
    run.font.name = FONT_BODY
    run.font.size = Pt(8.5)
    run.font.color.rgb = BODY
    if i < len(contact_items) - 1:
        run = p_contact.add_run("   ·   ")
        run.font.name = FONT_BODY
        run.font.size = Pt(8.5)
        run.font.color.rgb = LIGHT

add_thin_rule(doc)

# ── Profile ────────────────────────────────────────────────────────────
add_section_heading(doc, "Profile")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
add_rich_text(p, profile, Pt(9.5), BODY)

# ── Experience ─────────────────────────────────────────────────────────
add_section_heading(doc, "Experience")

for job in work:
    date_from = format_date(job["from"])
    date_to = format_date(job["to"]) if job["to"] else "Present"
    date_str = f"{date_from} – {date_to}"

    # Job title + company line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run(job["title"])
    run.font.name = FONT_HEAD
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    run.font.bold = True

    run = p.add_run(f"  —  {job['company']}")
    run.font.name = FONT_HEAD
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT
    run.font.bold = False

    # Meta line: location, type, dates
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    meta = f"{job['location']}  ·  {job['type']}  ·  {date_str}"
    run = p.add_run(meta)
    run.font.name = FONT_BODY
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT

    # Techs
    if job.get("techs"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(" · ".join(job["techs"]))
        run.font.name = FONT_BODY
        run.font.size = Pt(8)
        run.font.color.rgb = ACCENT
        run.font.italic = True

    # Description
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    add_rich_text(p, job["description"], Pt(9), BODY)


# ── Education ──────────────────────────────────────────────────────────
add_section_heading(doc, "Education")

for edu in education:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run(edu["diploma"])
    run.font.name = FONT_HEAD
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    run.font.bold = True

    run = p.add_run(f"  —  {edu['school']}")
    run.font.name = FONT_HEAD
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    meta = f"{edu['location']}  ·  {format_date(edu['date'])}"
    run = p.add_run(meta)
    run.font.name = FONT_BODY
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT

    if edu.get("description"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        add_rich_text(p, edu["description"], Pt(9), BODY)


# ── Skills ─────────────────────────────────────────────────────────────
add_section_heading(doc, "Skills")

# Create a table for skills: Name | Level | Experience | Related
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths
for cell in table.columns[0].cells:
    cell.width = Cm(3.0)
for cell in table.columns[1].cells:
    cell.width = Cm(2.5)
for cell in table.columns[2].cells:
    cell.width = Cm(2.5)
for cell in table.columns[3].cells:
    cell.width = Cm(9.0)

# Header row
hdr_cells = table.rows[0].cells
for i, text in enumerate(["Skill", "Level", "Experience", "Related Technologies"]):
    p = hdr_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = FONT_HEAD
    run.font.size = Pt(8.5)
    run.font.color.rgb = ACCENT
    run.font.bold = True
    # Bottom border for header
    set_cell_border(hdr_cells[i], bottom={"sz": 4, "color": "2B6CB3"})

# Data rows
for skill in skills:
    row_cells = table.add_row().cells

    p = row_cells[0].paragraphs[0]
    run = p.add_run(skill["name"])
    run.font.name = FONT_BODY
    run.font.size = Pt(9)
    run.font.color.rgb = DARK
    run.font.bold = True

    p = row_cells[1].paragraphs[0]
    run = p.add_run(skill["level"])
    run.font.name = FONT_BODY
    run.font.size = Pt(9)
    run.font.color.rgb = BODY

    p = row_cells[2].paragraphs[0]
    run = p.add_run(f"{skill['experience']} years")
    run.font.name = FONT_BODY
    run.font.size = Pt(9)
    run.font.color.rgb = BODY

    p = row_cells[3].paragraphs[0]
    run = p.add_run(", ".join(skill["related"]))
    run.font.name = FONT_BODY
    run.font.size = Pt(8.5)
    run.font.color.rgb = LIGHT

    # Light bottom border
    for cell in row_cells:
        set_cell_border(cell, bottom={"sz": 2, "color": "E0E0E0"})

# Remove all other table borders
tbl = table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'</w:tblBorders>'
)
tblPr.append(borders)


# ── Projects ───────────────────────────────────────────────────────────
add_section_heading(doc, "Personal Projects")

for proj in projects:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(proj["name"])
    run.font.name = FONT_HEAD
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    run.font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    add_rich_text(p, proj["description"], Pt(9), BODY)


# ── Save ───────────────────────────────────────────────────────────────
output_path = "Antoine_Jaussoin_CV.docx"
doc.save(output_path)
print(f"CV generated: {output_path}")
